import os
import hashlib
from typing import List

import streamlit as st
from dotenv import load_dotenv

from openai import OpenAI
from pinecone import Pinecone

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_pinecone import PineconeVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter

from pypdf import PdfReader
from docx import Document as DocxDocument


# ============================================================
# Configuration
# ============================================================

load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")

PINECONE_INDEX_NAME = os.getenv(
    "PINECONE_INDEX_NAME",
    "legal-documents",
)

EMBEDDING_MODEL = os.getenv(
    "EMBEDDING_MODEL",
    "nvidia/llama-nemotron-embed-vl-1b-v2:free",
)

EMBEDDING_DIMENSIONS = int(
    os.getenv("EMBEDDING_DIMENSIONS", "1024")
)

CHAT_MODEL = os.getenv(
    "CHAT_MODEL",
    "google/gemini-2.5-flash",
)


# ============================================================
# Validation
# ============================================================

if not OPENROUTER_API_KEY:
    st.error("OPENROUTER_API_KEY is missing from .env")
    st.stop()

if not PINECONE_API_KEY:
    st.error("PINECONE_API_KEY is missing from .env")
    st.stop()


# ============================================================
# Page Configuration
# ============================================================

st.set_page_config(
    page_title="Legal AI Assistant",
    page_icon="⚖️",
    layout="wide",
)


# ============================================================
# Custom Embedding Class
# ============================================================

class OpenRouterEmbeddings(Embeddings):

    def __init__(
        self,
        api_key: str,
        model: str = EMBEDDING_MODEL,
        dimensions: int = EMBEDDING_DIMENSIONS,
    ):
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
        )

        self.model = model
        self.dimensions = dimensions

    def embed_documents(
        self,
        texts: List[str],
    ) -> List[List[float]]:

        if not texts:
            return []

        response = self.client.embeddings.create(
            model=self.model,
            input=texts,
            dimensions=self.dimensions,
            encoding_format="float",
        )

        return [
            item.embedding
            for item in response.data
        ]

    def embed_query(
        self,
        text: str,
    ) -> List[float]:

        response = self.client.embeddings.create(
            model=self.model,
            input=text,
            dimensions=self.dimensions,
            encoding_format="float",
        )

        return response.data[0].embedding


# ============================================================
# Clients
# ============================================================

@st.cache_resource
def get_openrouter_client():

    return OpenAI(
        api_key=OPENROUTER_API_KEY,
        base_url="https://openrouter.ai/api/v1",
    )


@st.cache_resource
def get_embeddings():

    return OpenRouterEmbeddings(
        api_key=OPENROUTER_API_KEY,
    )


@st.cache_resource
def get_vector_store():

    pc = Pinecone(
        api_key=PINECONE_API_KEY,
    )

    index = pc.Index(
        PINECONE_INDEX_NAME
    )

    return PineconeVectorStore(
        index=index,
        embedding=get_embeddings(),
    )


client = get_openrouter_client()
vector_store = get_vector_store()


# ============================================================
# Document Extraction
# ============================================================

def extract_pdf(file) -> str:

    reader = PdfReader(file)

    pages = []

    for page in reader.pages:

        text = page.extract_text()

        if text:
            pages.append(text)

    return "\n\n".join(pages)


def extract_docx(file) -> str:

    doc = DocxDocument(file)

    paragraphs = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    ]

    return "\n\n".join(paragraphs)


def extract_txt(file) -> str:

    return file.getvalue().decode(
        "utf-8",
        errors="ignore",
    )


def extract_text(file) -> str:

    extension = file.name.lower().split(".")[-1]

    if extension == "pdf":
        return extract_pdf(file)

    if extension == "docx":
        return extract_docx(file)

    if extension == "txt":
        return extract_txt(file)

    raise ValueError(
        f"Unsupported file type: .{extension}"
    )


# ============================================================
# Hash File
# ============================================================

def get_file_hash(file) -> str:

    return hashlib.sha256(
        file.getvalue()
    ).hexdigest()


# ============================================================
# Chunking
# ============================================================

def create_chunks(
    text: str,
    file_name: str,
    file_hash: str,
) -> List[Document]:

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=200,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            "",
        ],
    )

    chunks = splitter.split_text(text)

    documents = []

    for index, chunk in enumerate(chunks):

        if not chunk.strip():
            continue

        documents.append(
            Document(
                page_content=chunk,
                metadata={
                    "source": file_name,
                    "file_hash": file_hash,
                    "chunk_id": index,
                },
            )
        )

    return documents


# ============================================================
# Document Ingestion
# ============================================================

def ingest_document(file):

    file_hash = get_file_hash(file)

    text = extract_text(file)

    if not text.strip():

        raise ValueError(
            "No readable text was found in the document."
        )

    documents = create_chunks(
        text=text,
        file_name=file.name,
        file_hash=file_hash,
    )

    if not documents:

        raise ValueError(
            "Document produced zero chunks."
        )

    # Deterministic IDs.
    # Uploading the same document again generates
    # the same IDs instead of creating random duplicates.
    ids = []

    for document in documents:

        chunk_hash = hashlib.sha256(
            document.page_content.encode("utf-8")
        ).hexdigest()

        vector_id = (
            f"{file_hash}-{chunk_hash}"
        )

        ids.append(vector_id)

    vector_store.add_documents(
        documents=documents,
        ids=ids,
    )

    return {
        "file_name": file.name,
        "chunks": len(documents),
        "file_hash": file_hash,
    }


# ============================================================
# Retrieve Documents
# ============================================================

def retrieve_documents(
    query: str,
    k: int = 5,
):

    results = vector_store.similarity_search_with_score(
        query=query,
        k=k,
    )

    return results


# ============================================================
# Generate Answer
# ============================================================

def generate_answer(
    query: str,
    results,
    chat_history,
):

    if not results:

        return (
            "I don't know based on the provided documents.",
            [],
        )

    # Build context
    context_parts = []

    for i, (doc, score) in enumerate(
        results,
        start=1,
    ):

        source = doc.metadata.get(
            "source",
            "Unknown",
        )

        context_parts.append(
            f"""
SOURCE {i}
File: {source}
Similarity Score: {score:.4f}

Content:
{doc.page_content}
"""
        )

    context = "\n\n".join(
        context_parts
    )

    # Keep conversation reasonably small
    previous_messages = ""

    for message in chat_history[-6:]:

        previous_messages += (
            f"{message['role'].upper()}: "
            f"{message['content']}\n"
        )

    prompt = f"""
You are an expert legal AI assistant.

Your job is to answer questions using ONLY
the information contained in the provided documents.

Rules:

1. Do not use outside knowledge.
2. Do not invent legal facts.
3. If the answer is not supported by the documents,
   respond exactly:

"I don't know based on the provided documents."

4. When possible, mention the source filename.
5. Give a clear and concise answer.
6. Treat previous conversation as context,
   but documents are the source of truth.

Previous conversation:
{previous_messages}

Retrieved documents:
{context}

User question:
{query}

Answer:
"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a document-grounded "
                    "legal AI assistant."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0,
        max_tokens=800,
    )

    answer = response.choices[0].message.content

    return answer, results


# ============================================================
# Session State
# ============================================================

if "messages" not in st.session_state:

    st.session_state.messages = []


# ============================================================
# Sidebar
# ============================================================

with st.sidebar:

    st.title("Document Management")

    st.markdown(
        "Upload legal documents and add them "
        "to the Pinecone knowledge base."
    )

    uploaded_files = st.file_uploader(
        "Upload documents",
        type=[
            "pdf",
            "docx",
            "txt",
        ],
        accept_multiple_files=True,
    )

    if uploaded_files:

        if st.button(
            "Process Documents",
            type="primary",
            use_container_width=True,
        ):

            progress = st.progress(0)

            total = len(uploaded_files)

            for index, file in enumerate(
                uploaded_files
            ):

                try:

                    result = ingest_document(file)

                    st.success(
                        f"Added {result['file_name']} "
                        f"({result['chunks']} chunks)"
                    )

                except Exception as e:

                    st.error(
                        f"Failed to process "
                        f"{file.name}: {e}"
                    )

                progress.progress(
                    (index + 1) / total
                )

    st.divider()

    st.subheader("Retrieval")

    retrieval_k = st.slider(
        "Documents to retrieve",
        min_value=1,
        max_value=10,
        value=5,
    )

    st.divider()

    if st.button(
        "Clear Conversation",
        use_container_width=True,
    ):

        st.session_state.messages = []

        st.rerun()


# ============================================================
# Main UI
# ============================================================

st.title("⚖️ Legal AI Assistant")

st.caption(
    "Chat with your uploaded legal documents "
    "using OpenRouter + Pinecone + RAG."
)


# ============================================================
# Display Chat History
# ============================================================

for message in st.session_state.messages:

    with st.chat_message(
        message["role"]
    ):

        st.markdown(
            message["content"]
        )

        # Display sources for assistant messages
        if (
            message["role"] == "assistant"
            and message.get("sources")
        ):

            with st.expander(
                "View retrieved sources"
            ):

                for source in message["sources"]:

                    st.markdown(
                        f"""
**File:** {source["file"]}

**Score:** {source["score"]:.4f}

**Content:**

{source["content"]}
"""
                    )


# ============================================================
# Chat Input
# ============================================================

query = st.chat_input(
    "Ask a question about your legal documents..."
)


if query:

    # User message
    st.session_state.messages.append(
        {
            "role": "user",
            "content": query,
        }
    )

    with st.chat_message("user"):
        st.markdown(query)

    with st.chat_message("assistant"):

        with st.spinner(
            "Searching documents..."
        ):

            try:

                results = retrieve_documents(
                    query=query,
                    k=retrieval_k,
                )

                answer, sources = generate_answer(
                    query=query,
                    results=results,
                    chat_history=st.session_state.messages[:-1],
                )

                st.markdown(answer)

                # Prepare source data
                source_data = []

                for doc, score in results:

                    source_data.append(
                        {
                            "file": doc.metadata.get(
                                "source",
                                "Unknown",
                            ),
                            "score": score,
                            "content": doc.page_content,
                        }
                    )

                # Show sources
                if source_data:

                    with st.expander(
                        "View retrieved sources"
                    ):

                        for source in source_data:

                            st.markdown(
                                f"""
**File:** {source["file"]}

**Similarity Score:** {source["score"]:.4f}

{source["content"]}
"""
                            )

                # Save assistant response
                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": answer,
                        "sources": source_data,
                    }
                )

            except Exception as e:

                error_message = (
                    f"An error occurred: "
                    f"{type(e).__name__}: {e}"
                )

                st.error(error_message)

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": error_message,
                        "sources": [],
                    }
                )